import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def generate_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('AI Energy Optimizer: Frontend Dashboard Architecture & Workflow', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 1. Introduction
    doc.add_heading('1. Overview of the Application', level=1)
    doc.add_paragraph(
        "The AI Energy Optimizer Dashboard is a real-time React application that interfaces with a FastAPI machine learning backend. "
        "It provides a visually intuitive interface for users to submit hourly energy consumption readings. Behind the scenes, the dashboard "
        "receives a comprehensive payload containing future predictions (LightGBM), anomaly risks (Isolation Forest), and cost-saving schedules (Linear Programming)."
    )

    # 2. General Workflow: Submitting a Reading
    doc.add_heading('2. What Happens When You Submit a Reading?', level=1)
    doc.add_paragraph(
        "When the user selects a Date/Time and enters a consumption value (kWh) in the Prediction Form, the following events occur:"
    )
    process_points = [
        "The Dashboard component triggers an HTTP POST request to the backend's `/predict_full` endpoint.",
        "The backend processes the reading, updates the active rolling memory buffer, and runs three separate AI operations: forecasting the next step, detecting anomalies, and optimizing the schedule.",
        "The backend returns a unified JSON dictionary containing 'forecast', 'anomaly', and 'optimization' keys.",
        "The React state is immediately updated, pushing the new data object into a historical memory array limit of 24 objects (MAX_HISTORY).",
        "This state update propagates downward, prompting all sub-components (charts, UI panels, anomaly indicators) to visually re-render with the new data."
    ]
    for p in process_points:
        doc.add_paragraph(p, style='List Bullet')

    # 3. Handling High Readings / Anomalies
    doc.add_heading('3. What Happens If the Reading is High?', level=1)
    doc.add_paragraph(
        "A reading is classified as 'high' or 'anomalous' when it breaches both biological statistical limits (e.g. 3.5 standard deviations above the 24-hour moving average) and the boundaries set by the unsupervised Isolation Forest."
    )
    doc.add_paragraph(
        "In the Dashboard, the AnomalyIndicator component reacts in real-time:"
    )
    anomaly_points = [
        "Status Badge: Transitions from a green 'Normal' pill to a red 'Anomaly' or yellow 'Warning' badge.",
        "Internal Metrics: The component maps out exactly why it fired. It displays the safe Mean, the Sigma Distance, the Load Ratio multiplier, and flags whether the Statistical Gate or Isolation Forest caught the spike.",
        "Timeline Viz: The component accesses the last 5 readings of the user's session and displays a textual timeline. The anomalous hour will be boldly highlighted in red with an 'ALERT' tag."
    ]
    for a in anomaly_points:
        doc.add_paragraph(a, style='List Bullet')

    # 4. Interactive Graphs and Features
    doc.add_heading('4. How the Line Graph is Drawn (ForecastChart)', level=1)
    doc.add_paragraph(
        "The primary line graph is built using Recharts. It plots the 'history' array maintained globally by the main Dashboard component."
    )
    doc.add_paragraph(
        "What it shows and indicates:", style='List Bullet'
    )
    doc.add_paragraph(
        "Actual Consumption Trace: Tracks the exact history of readings the user has verified. Plotted as a solid line.", style='List Bullet 2'
    )
    doc.add_paragraph(
        "Forecast Trace: Tracks what the LightGBM algorithm predicted for that hour before it happened.", style='List Bullet 2'
    )
    doc.add_paragraph(
        "Shaded Confidence Interval: Using an <Area> path layer, it draws a shaded boundary using the 5th and 95th quantile predictions from the XGBoost models. This represents the 'safe' AI expectation zone.", style='List Bullet 2'
    )
    doc.add_paragraph(
        "Anomaly Dots: If a specific hour fired the anomaly flag, the chart overlays a red dot exactly on top of the intersection point.", style='List Bullet 2'
    )

    # 5. LP Cost Optimizer
    doc.add_heading('5. The Linear Programming (LP) Cost Optimizer', level=1)
    doc.add_paragraph(
        "The LP Optimizer is a mathematical solver running continuously behind the scenes. When a reading is evaluated, the backend "
        "projects the user's trajectory out over the next 48 hours."
    )
    doc.add_paragraph(
        "What it does: It applies a constraint-based shift using the Highs Solver. It takes up to 20% of the energy consumed during peak hours (where rates are expensive, Rs. 9.0/kWh) and forcibly shifts it to off-peak night cycles (Rs. 3.5/kWh)."
    )
    
    doc.add_heading('What the Bar Graph Denotes (OptimizationPanel)', level=2)
    doc.add_paragraph(
        "The Optimization Panel renders a simple UI summary comparing the financial impact:"
    )
    bar_features = [
        "It uses a dual-BarChart rendering two bars side-by-side representing total cost over the optimization horizon.",
        "Red/Purple Bar: Denotes the 'Original Cost' - exactly what the user would pay on their current trajectory without optimization.",
        "Green Bar: Denotes the 'Optimized Cost' - the new, cheaper invoice total if exactly 20% of peak loads are delayed.",
        "Financial Statistics: The UI specifically highlights total Rupees saved and the percentage yield of the cut."
    ]
    for b in bar_features:
        doc.add_paragraph(b, style='List Bullet')

    # 6. React Component Workflow Breakdown
    doc.add_heading('6. Component Workflow Lifecycle', level=1)
    
    components = {
        "Dashboard.jsx": "The 'Brain'. Stores global state variables (loading flags, the history array, the latest JSON output) and handles all API calls so child components remain visually pure.",
        "PredictionForm.jsx": "The 'Input'. Traps the User's form submissions, prevents default page reloading, and triggers the `onSubmit` callback sent down mechanically from the Dashboard.",
        "ForecastChart.jsx": "The 'Chronicle'. Absorbs the Dashboard's history array, scaling its X/Y axis automatically, mapping confidence areas organically.",
        "AnomalyIndicator.jsx": "The 'Alarm'. Reads the specific `anomaly` fragment of the response payload, turning parameters into color-coded blocks.",
        "OptimizationPanel.jsx": "The 'Bank'. Unpacks the `optimization` JSON fragment, drawing the two budget bars and formatting integers to currency strings."
    }

    for comp, desc in components.items():
        p = doc.add_paragraph()
        run = p.add_run(f"• {comp}: ")
        run.bold = True
        p.add_run(desc)

    # Output path
    output_tgt = os.path.join(os.getcwd(), 'React_Dashboard_Technical_Documentation.docx')
    doc.save(output_tgt)
    print(f"Documentation generated at: {output_tgt}")

if __name__ == "__main__":
    generate_report()
